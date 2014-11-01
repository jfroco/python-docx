# encoding: utf-8

"""
Step implementations for table-related features
"""

from __future__ import absolute_import, print_function, unicode_literals

from behave import given, then, when

from docx import Document
from docx.shared import Inches
from docx.table import (
    _Cell, _Column, _ColumnCells, _Columns, _Row, _RowCells, _Rows
)

from helpers import test_docx


# given ===================================================

@given('a 2 x 2 table')
def given_a_2x2_table(context):
    context.table_ = Document().add_table(rows=2, cols=2)


@given('a 3x3 table containing {span_state}')
def given_a_3x3_table_containing_span_state(context, span_state):
    table_idx = {
        'only uniform cells': 0,
        'a horizontal span':  1,
        'a vertical span':    2,
        'a combined span':    3,
    }[span_state]
    document = Document(test_docx('tbl-collections'))
    context.table_ = document.tables[table_idx]


@given('a column cell collection having two cells')
def given_a_column_cell_collection_having_two_cells(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.cells = document.tables[0].columns[0].cells


@given('a column collection having two columns')
def given_a_column_collection_having_two_columns(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.columns = document.tables[0].columns


@given('a row cell collection having two cells')
def given_a_row_cell_collection_having_two_cells(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.cells = document.tables[0].rows[0].cells


@given('a row collection having two rows')
def given_a_row_collection_having_two_rows(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.rows = document.tables[0].rows


@given('a table')
def given_a_table(context):
    context.table_ = Document().add_table(rows=2, cols=2)


@given('a table cell having a width of {width}')
def given_a_table_cell_having_a_width_of_width(context, width):
    table_idx = {'no explicit setting': 0, '1 inch': 1, '2 inches': 2}[width]
    document = Document(test_docx('tbl-props'))
    table = document.tables[table_idx]
    cell = table.cell(0, 0)
    context.cell = cell


@given('a table column having a width of {width_desc}')
def given_a_table_having_a_width_of_width_desc(context, width_desc):
    col_idx = {
        'no explicit setting': 0,
        '1440':                1,
    }[width_desc]
    docx_path = test_docx('tbl-col-props')
    document = Document(docx_path)
    context.column = document.tables[0].columns[col_idx]


@given('a table having an applied style')
def given_a_table_having_an_applied_style(context):
    docx_path = test_docx('tbl-having-applied-style')
    document = Document(docx_path)
    context.table_ = document.tables[0]


@given('a table having an autofit layout of {autofit}')
def given_a_table_having_an_autofit_layout_of_autofit(context, autofit):
    tbl_idx = {
        'no explicit setting': 0,
        'autofit':             1,
        'fixed':               2,
    }[autofit]
    document = Document(test_docx('tbl-props'))
    context.table_ = document.tables[tbl_idx]


@given('a table having two columns')
def given_a_table_having_two_columns(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    # context.table is used internally by behave, underscore added
    # to distinguish this one
    context.table_ = document.tables[0]


@given('a table having two rows')
def given_a_table_having_two_rows(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.table_ = document.tables[0]


@given('a table column having two cells')
def given_a_table_column_having_two_cells(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.column = document.tables[0].columns[0]


@given('a table row having two cells')
def given_a_table_row_having_two_cells(context):
    docx_path = test_docx('blk-containing-table')
    document = Document(docx_path)
    context.row = document.tables[0].rows[0]


# when =====================================================

@when('I add a column to the table')
def when_add_column_to_table(context):
    table = context.table_
    context.column = table.add_column()


@when('I add a row to the table')
def when_add_row_to_table(context):
    table = context.table_
    context.row = table.add_row()


@when('I apply a style to the table')
def when_apply_style_to_table(context):
    table = context.table_
    table.style = 'LightShading-Accent1'


@when('I set the cell width to {width}')
def when_I_set_the_cell_width_to_width(context, width):
    new_value = {'1 inch': Inches(1)}[width]
    context.cell.width = new_value


@when('I set the column width to {width_emu}')
def when_I_set_the_column_width_to_width_emu(context, width_emu):
    new_value = None if width_emu == 'None' else int(width_emu)
    context.column.width = new_value


@when('I set the table autofit to {setting}')
def when_I_set_the_table_autofit_to_setting(context, setting):
    new_value = {'autofit': True, 'fixed': False}[setting]
    table = context.table_
    table.autofit = new_value


# then =====================================================

@then('I can access a cell using its row and column indices')
def then_can_access_cell_using_its_row_and_col_indices(context):
    table = context.table_
    for row_idx in range(2):
        for col_idx in range(2):
            cell = table.cell(row_idx, col_idx)
            assert isinstance(cell, _Cell)


@then('I can access a collection column by index')
def then_can_access_collection_column_by_index(context):
    columns = context.columns
    for idx in range(2):
        column = columns[idx]
        assert isinstance(column, _Column)


@then('I can access a collection row by index')
def then_can_access_collection_row_by_index(context):
    rows = context.rows
    for idx in range(2):
        row = rows[idx]
        assert isinstance(row, _Row)


@then('I can access a column cell by index')
def then_can_access_column_cell_by_index(context):
    cells = context.cells
    for idx in range(2):
        cell = cells[idx]
        assert isinstance(cell, _Cell)


@then('I can access a row cell by index')
def then_can_access_row_cell_by_index(context):
    cells = context.cells
    for idx in range(2):
        cell = cells[idx]
        assert isinstance(cell, _Cell)


@then('I can access the cell collection of the column')
def then_can_access_cell_collection_of_column(context):
    column = context.column
    cells = column.cells
    assert isinstance(cells, _ColumnCells)


@then('I can access the cell collection of the row')
def then_can_access_cell_collection_of_row(context):
    row = context.row
    cells = row.cells
    assert isinstance(cells, _RowCells)


@then('I can access the column collection of the table')
def then_can_access_column_collection_of_table(context):
    table = context.table_
    columns = table.columns
    assert isinstance(columns, _Columns)


@then('I can access the row collection of the table')
def then_can_access_row_collection_of_table(context):
    table = context.table_
    rows = table.rows
    assert isinstance(rows, _Rows)


@then('I can get the length of the column cell collection')
def then_can_get_length_of_column_cell_collection(context):
    column = context.column
    cells = column.cells
    assert len(cells) == 2


@then('I can get the length of the row cell collection')
def then_can_get_length_of_row_cell_collection(context):
    row = context.row
    cells = row.cells
    assert len(cells) == 2


@then('I can get the table style name')
def then_can_get_table_style_name(context):
    table = context.table_
    msg = "got '%s'" % table.style
    assert table.style == 'LightShading-Accent1', msg


@then('I can iterate over the column cells')
def then_can_iterate_over_the_column_cells(context):
    cells = context.cells
    actual_count = 0
    for cell in cells:
        actual_count += 1
        assert isinstance(cell, _Cell)
    assert actual_count == 2


@then('I can iterate over the column collection')
def then_can_iterate_over_column_collection(context):
    columns = context.columns
    actual_count = 0
    for column in columns:
        actual_count += 1
        assert isinstance(column, _Column)
    assert actual_count == 2


@then('I can iterate over the row cells')
def then_can_iterate_over_the_row_cells(context):
    cells = context.cells
    actual_count = 0
    for cell in cells:
        actual_count += 1
        assert isinstance(cell, _Cell)
    assert actual_count == 2


@then('I can iterate over the row collection')
def then_can_iterate_over_row_collection(context):
    rows = context.rows
    actual_count = 0
    for row in rows:
        actual_count += 1
        assert isinstance(row, _Row)
    assert actual_count == 2


@then('len(column.cells) is {rows_str} for each of its columns')
def then_len_column_cells_is_rows_for_each_of_its_columns(context, rows_str):
    table = context.table_
    rows = int(rows_str)
    assert len(table.columns) > 0
    for column in table.columns:
        assert len(column.cells) == rows, 'got %s' % len(column.cells)


@then('len(row.cells) is {cols_str} for each of its rows')
def then_len_row_cells_is_cols_for_each_of_its_rows(context, cols_str):
    table = context.table_
    cols = int(cols_str)
    assert len(table.rows) > 0
    for row in table.rows:
        assert len(row.cells) == cols, 'got %s' % len(row.cells)


@then('len(table.columns) is {cols_str}')
def then_len_table_columns_is_cols(context, cols_str):
    table = context.table_
    cols = int(cols_str)
    assert len(table.columns) == cols, 'got %s' % len(table.columns)


@then('len(table.rows) is {rows_str}')
def then_len_table_rows_is_rows(context, rows_str):
    table = context.table_
    rows = int(rows_str)
    assert len(table.rows) == rows, 'got %s' % len(table.rows)


@then('the length of the column collection is 2')
def then_len_of_column_collection_is_2(context):
    columns = context.table_.columns
    assert len(columns) == 2


@then('the length of the row collection is 2')
def then_len_of_row_collection_is_2(context):
    rows = context.table_.rows
    assert len(rows) == 2


@then('the new column has 2 cells')
def then_new_column_has_2_cells(context):
    assert len(context.column.cells) == 2


@then('the new row has 2 cells')
def then_new_row_has_2_cells(context):
    assert len(context.row.cells) == 2


@then('the reported autofit setting is {autofit}')
def then_the_reported_autofit_setting_is_autofit(context, autofit):
    expected_value = {'autofit': True, 'fixed': False}[autofit]
    table = context.table_
    assert table.autofit is expected_value


@then('the reported column width is {width_emu}')
def then_the_reported_column_width_is_width_emu(context, width_emu):
    expected_value = None if width_emu == 'None' else int(width_emu)
    assert context.column.width == expected_value, (
        'got %s' % context.column.width
    )


@then('the reported width of the cell is {width}')
def then_the_reported_width_of_the_cell_is_width(context, width):
    expected_width = {'None': None, '1 inch': Inches(1)}[width]
    actual_width = context.cell.width
    assert actual_width == expected_width, (
        'expected %s, got %s' % (expected_width, actual_width)
    )


@then('the table style matches the name I applied')
def then_table_style_matches_name_applied(context):
    table = context.table_
    tmpl = "table.style doesn't match, got '%s'"
    assert table.style == 'LightShading-Accent1', tmpl % table.style


@then('the table has {count} columns')
def then_table_has_count_columns(context, count):
    column_count = int(count)
    columns = context.table_.columns
    assert len(columns) == column_count


@then('the table has {count} rows')
def then_table_has_count_rows(context, count):
    row_count = int(count)
    rows = context.table_.rows
    assert len(rows) == row_count
